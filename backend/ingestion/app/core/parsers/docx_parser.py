"""
DOCX Parser — utilise python-docx pour extraire texte, tableaux et métadonnées
des fichiers Word (.docx).
"""
from docx import Document
from docx.oxml.ns import qn
from .base import BaseParser, ParseResult


class DocxParser(BaseParser):

    def supported_extensions(self) -> list[str]:
        return [".docx", ".doc"]

    def parse(self, file_path: str) -> ParseResult:
        doc = Document(file_path)

        # ── Paragraphes ───────────────────────────────────────────────────────
        paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]

        # ── Tableaux ──────────────────────────────────────────────────────────
        tables = []
        for table in doc.tables:
            rows = []
            for row in table.rows:
                rows.append([cell.text.strip() for cell in row.cells])
            tables.append(rows)

        # ── Métadonnées ───────────────────────────────────────────────────────
        props = doc.core_properties
        metadata = {
            "title":    props.title or "",
            "author":   props.author or "",
            "subject":  props.subject or "",
            "created":  str(props.created) if props.created else "",
            "modified": str(props.modified) if props.modified else "",
        }

        full_text = "\n\n".join(paragraphs)

        return ParseResult(
            text=full_text,
            metadata=metadata,
            pages=[full_text],
            tables=tables,
            images_text=[],
        )
